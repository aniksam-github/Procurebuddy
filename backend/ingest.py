import os
import shutil
import stat
from pathlib import Path

import pdfplumber
import pytesseract
from docx import Document as DocxDocument
from dotenv import load_dotenv
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pdf2image import convert_from_path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
ENV_FILE = PROJECT_ROOT / ".env"
DATA_PATH = PROJECT_ROOT / "data"
DB_PATH = PROJECT_ROOT / "chroma_db"
EMBEDDING_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
SUPPORTED_DOC_EXTENSIONS = {".pdf", ".docx", ".txt"}

load_dotenv(dotenv_path=ENV_FILE)


def _remove_readonly(func, path, _exc_info):
    os.chmod(path, stat.S_IWRITE)
    func(path)


def extract_metadata_from_filename(filename: str):
    lowered = filename.lower()
    doc_type = "Manual"
    year = "2019"

    if "special provisions" in lowered or "om" in lowered or "amendment" in lowered or "preference" in lowered:
        doc_type = "Office Memorandum / Amendment"
        year = "2025"
    elif "gfr" in lowered:
        doc_type = "GFR Rule"
        year = "2017"

    return doc_type, year


def _base_metadata(path: Path) -> dict:
    doc_type, year = extract_metadata_from_filename(path.name)
    return {
        "source": str(path),
        "source_file": path.name,
        "doc_type": doc_type,
        "year": year,
    }


def _ocr_pdf_page(path: Path, page_number: int) -> str:
    try:
        images = convert_from_path(str(path), first_page=page_number, last_page=page_number)
        if not images:
            return ""
        return pytesseract.image_to_string(images[0]).strip()
    except Exception:
        return ""


def _load_pdf(path: Path) -> tuple[list[Document], int]:
    documents = []
    ocr_pages = 0

    with pdfplumber.open(str(path)) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            text = (page.extract_text() or "").strip()
            used_ocr = False

            if not text:
                text = _ocr_pdf_page(path, index)
                used_ocr = bool(text)

            if not text:
                continue

            metadata = {
                **_base_metadata(path),
                "page": index,
                "ocr_used": used_ocr,
            }
            documents.append(Document(page_content=text, metadata=metadata))
            if used_ocr:
                ocr_pages += 1

    return documents, ocr_pages


def _load_docx(path: Path) -> tuple[list[Document], int]:
    doc = DocxDocument(str(path))
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()).strip()
    if not text:
        return [], 0
    return [Document(page_content=text, metadata=_base_metadata(path))], 0


def _load_txt(path: Path) -> tuple[list[Document], int]:
    text = path.read_text(encoding="utf-8", errors="ignore").strip()
    if not text:
        return [], 0
    return [Document(page_content=text, metadata=_base_metadata(path))], 0


def _load_supported_documents():
    documents = []
    ocr_pages = 0
    processed_files = []

    for path in sorted(DATA_PATH.iterdir(), key=lambda item: item.name.lower()):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_DOC_EXTENSIONS:
            continue

        if path.suffix.lower() == ".pdf":
            loaded, ocr_count = _load_pdf(path)
        elif path.suffix.lower() == ".docx":
            loaded, ocr_count = _load_docx(path)
        else:
            loaded, ocr_count = _load_txt(path)

        if loaded:
            documents.extend(loaded)
            processed_files.append(path.name)
        ocr_pages += ocr_count

    return documents, processed_files, ocr_pages


def create_vector_db():
    if not DATA_PATH.exists():
        raise FileNotFoundError(f"Data directory not found: {DATA_PATH}")

    documents, processed_files, ocr_pages = _load_supported_documents()
    if not documents:
        raise FileNotFoundError(f"No supported documents with extractable text found in {DATA_PATH}")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=150,
        separators=["\n\n", "\n", ".", " "],
    )
    chunks = splitter.split_documents(documents)

    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)
    if DB_PATH.exists():
        shutil.rmtree(DB_PATH, onerror=_remove_readonly)
    Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        persist_directory=str(DB_PATH),
    )

    return {
        "vector_db_path": str(DB_PATH),
        "document_count": len(processed_files),
        "chunk_count": len(chunks),
        "ocr_pages": ocr_pages,
        "files": processed_files,
    }
